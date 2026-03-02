"""
Parse Mauritanian recipes from Word document.
The document has corrupted/overlapping text, so we need to extract carefully.
"""

from docx import Document
from docx.oxml.ns import qn
import json
import re

def extract_text_from_docx(filepath):
    """Extract all text including from tables and text boxes."""
    doc = Document(filepath)
    
    all_text = []
    
    # Get text from paragraphs
    for para in doc.paragraphs:
        text = para.text.strip()
        if text:
            all_text.append(text)
    
    # Get text from tables
    for table in doc.tables:
        for row in table.rows:
            for cell in row.cells:
                text = cell.text.strip()
                if text:
                    all_text.append(text)
    
    return all_text

def parse_recipes_from_text(text_lines):
    """Try to parse recipe JSON objects from the text."""
    # Join all lines
    full_text = '\n'.join(text_lines)
    
    # Try to find JSON-like patterns for recipe names
    # Pattern to find recipe name declarations
    name_pattern = r'"name":\s*"([^"]+)"'
    names = re.findall(name_pattern, full_text)
    
    print(f"Found {len(names)} recipe names:")
    for i, name in enumerate(names, 1):
        print(f"  {i}. {name}")
    
    return names

def main():
    filepath = r'C:\Users\Admin\Downloads\Document (4) (1).docx'
    
    print("=" * 60)
    print("PARSING MAURITANIAN RECIPES DOCUMENT")
    print("=" * 60)
    
    # Extract text
    text_lines = extract_text_from_docx(filepath)
    
    print(f"\nExtracted {len(text_lines)} text blocks")
    print("\n--- RAW TEXT ---")
    for i, line in enumerate(text_lines[:50]):  # First 50 lines
        print(f"{i+1}: {line[:100]}{'...' if len(line) > 100 else ''}")
    
    print("\n--- RECIPE NAMES FOUND ---")
    names = parse_recipes_from_text(text_lines)

if __name__ == "__main__":
    main()
